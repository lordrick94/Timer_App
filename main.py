import os
import sys
import math

import xlwings as xw
from kivy.config import Config
from docx import *
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, Inches
from docx.oxml.xmlchemy import OxmlElement
from docx.oxml.shared import qn

from kivy.app import App
from kivy.lang import Builder
from kivy.properties import StringProperty
from kivy.uix.tabbedpanel import TabbedPanel
from datetime import date
import kivy


def resource_path(relative_path):
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(os.path.abspath("."), relative_path)


kivy.require('1.9.0')

Config.set('graphics', 'width', '1200')
Config.set('graphics', 'height', '600')


class MainLayout(TabbedPanel):
    unit = StringProperty("seconds")

    def on_click_seconds(self, instance, value, units):
        if value == True:
            self.unit = units

    def on_click_minutes(self, instance, value, units):
        if value == True:
            self.unit = units

    certificate_number = StringProperty("34552")

    def cert_numb(self, widget):
        self.certificate_number = widget.text

    Equipment = StringProperty("TIMER")

    def equipment(self, widget):
        self.Equipment = widget.text

    Identification_Number = StringProperty("TBS/LAB/01")

    def identification_number(self, widget):
        self.Identification_Number = widget.text

    Manufacturer = StringProperty("Germany")

    def manufacturer(self, widget):
        self.Manufacturer = widget.text

    Serial_Number = StringProperty("99999")

    def serial_number(self, widget):
        self.Serial_Number = widget.text

    Readability = StringProperty("1s")

    def readability(self, widget):
        self.Readability = widget.text

    Location = StringProperty("TBS Time and Frequency Lab Ubungo")

    def location(self, widget):
        self.Location = widget.text

    Job_num = StringProperty("2021 - 999")

    def job_num(self, widget):
        self.Job_num = widget.text

    d1 = str(date.today().strftime("%Y-%m-%d"))
    Calibration_date = StringProperty(d1)

    def calibration_date(self, widget):
        self.Calibration_date = widget.text

    Issue_date = StringProperty(d1)

    def issue_date(self, widget):
        self.Issue_date = widget.text

    Customer_Name = StringProperty("Lordrick Apps")

    def customer_name(self, widget):
        self.Customer_Name = widget.text

    Postal_Address = StringProperty("7713")

    def postal_address(self, widget):
        self.Postal_Address = widget.text

    Region = StringProperty("Dodoma")

    def region(self, widget):
        self.Region = widget.text

    Country = StringProperty("Tanzania")

    def country(self, widget):
        self.Country = widget.text

    Metrologist = StringProperty("J. Yarrot")

    def metrologist(self, widget):
        self.Metrologist = widget.text

    Technical_Signatory = StringProperty("K. Shemhilu")

    def technical_signatory(self, widget):
        self.Technical_Signatory = widget.text

    Head_of_department = StringProperty("J. Mahilla")

    def head_of_department(self, widget):
        self.Head_of_department = widget.text

    Std_timer_id = StringProperty("TBS 8-381")

    def std_timer_id(self, widget):
        self.Std_timer_id = widget.text

    Std_timer_cert = StringProperty("58469")

    def std_timer_cert(self, widget):
        self.Std_timer_cert = widget.text

    Thermo_serial = StringProperty("MAC:98:8B:AD:20:C2:9F")

    def thermo_serial(self, widget):
        self.Thermo_serial = widget.text

    Thermo_cert = StringProperty("45326")

    def thermo_cert(self, widget):
        self.Thermo_cert = widget.text

    Valid = StringProperty("Dec -2021")

    def valid(self, widget):
        self.Valid = widget.text

    Method_of_calibration = StringProperty("MET - TF - 03")

    def method_of_calibration(self, widget):
        self.Method_of_calibration = widget.text

    Ambient_temperature = StringProperty("22,46")

    def ambient_temperature(self, widget):
        self.Ambient_temperature = widget.text

    Relative_humidity = StringProperty("50,0")

    def relative_humidity(self, widget):
        self.Relative_humidity = widget.text

    Resolution = StringProperty("0.5")

    def resolution(self, widget):
        self.Resolution = widget.text

    Remarks = StringProperty("DUT Means device under test")

    def remarks(self, widget):
        self.Remarks = widget.text

    D_1 = StringProperty("60")

    def d_1(self, widget):
        self.D_1 = widget.text

    D_2 = StringProperty("300")

    def d_2(self, widget):
        self.D_2 = widget.text

    D_3 = StringProperty("600")

    def d_3(self, widget):
        self.D_3 = widget.text

    D_4 = StringProperty("900")

    def d_4(self, widget):
        self.D_4 = widget.text

    D_5 = StringProperty("1200")

    def d_5(self, widget):
        self.D_5 = widget.text

    S_1 = StringProperty("61")

    def s_1(self, widget):
        self.S_1 = widget.text

    S_2 = StringProperty("301")

    def s_2(self, widget):
        self.S_2 = widget.text

    S_3 = StringProperty("600")

    def s_3(self, widget):
        self.S_3 = widget.text

    S_4 = StringProperty("899")

    def s_4(self, widget):
        self.S_4 = widget.text

    S_5 = StringProperty("1199")

    def s_5(self, widget):
        self.S_5 = widget.text

    def certificate_generator(self):
        path = "Certificates/"

        wb = xw.Book(resource_path(path + "Template.xlsx"))

        ws = wb.sheets['Sheet1']

        ws.range('B9').value = self.Customer_Name

        ws.range('B10').value = self.Equipment

        ws.range('B11').value = self.Identification_Number

        ws.range('B12').value = self.certificate_number

        ws.range('B13').value = self.Job_num

        ws.range('B14').value = self.Metrologist

        if self.unit == 'minutes':
            ws.range('D16').value = "Time interval (minutes)"
            ws.range('D18').value = ws.range('D18').value / 60
            ws.range('D19').value = ws.range('D19').value / 60
            ws.range('D20').value = ws.range('D20').value / 60
            ws.range('D22').value = ws.range('D22').value / 60

        res = float(self.Resolution)

        ws.range('D21').value = res

        uncertainty = ws.range('H27').value

        if 0.1 < uncertainty < 1:
            uncertainty = 1

        elif 0.01 < uncertainty < 0.1:
            uncertainty = 0.1

        else:
            uncertainty = math.ceil(uncertainty)

        ws.range('H28').value = f'U= ± {uncertainty}'

        wb.save(resource_path(path + f"{self.certificate_number}.xlsx"))

        doc = Document()

        sections = doc.sections

        for section in sections:
            section.top_margin = Inches(0.71)
            section.bottom_margin = Inches(0.71)
            section.left_margin = Inches(0.71)
            section.right_margin = Inches(0.55)

        sec_pr = doc.sections[0]._sectPr  # get the section properties el
        # create new borders el
        pg_borders = OxmlElement('w:pgBorders')
        # specifies how the relative positioning of the borders should be calculated
        pg_borders.set(qn('w:offsetFrom'), 'page')
        for border_name in ('top', 'left', 'bottom', 'right',):  # set all borders
            border_el = OxmlElement(f'w:{border_name}')
            border_el.set(qn('w:val'), 'double')  # a double line
            border_el.set(qn('w:sz'), '20')  # for meaning of remaining attrs please look docs
            border_el.set(qn('w:space'), '24')
            border_el.set(qn('w:color'), 'auto')
            pg_borders.append(border_el)  # register single border to border el
        sec_pr.append(pg_borders)  # apply border changes to section

        Certificate_No = self.certificate_number

        header = doc.sections[0].header
        para = header.add_paragraph()
        para.paragraph_format.space_after = 0
        para.paragraph_format.space_before = 0
        r = para.add_run()
        r.add_picture(resource_path("tbs logo.png"))
        r_space_1 = para.add_run("\t\t\t\t\t\t\t\t")
        r_1 = para.add_run()
        r_1.add_picture(resource_path("sadcas.png"))

        para_h = header.add_paragraph(f'\t\t\t\t\tCERTIFICATE NO: {Certificate_No}')
        para_h.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        para_h.style = doc.styles.add_style('Style name', WD_STYLE_TYPE.PARAGRAPH)
        font = para_h.style.font
        font.name = 'Arial'
        font.size = Pt(12)
        font.bold = True

        r_h = para_h.add_run("\t\t\t\t  Page 1 of 3")
        font = r_h.font
        font.name = 'Arial'
        font.size = Pt(12)
        font.bold = False

        head1 = doc.add_paragraph("TANZANIA BUREAU OF STANDARDS")
        head1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        head1.style = doc.styles.add_style('Style name_header_1', WD_STYLE_TYPE.PARAGRAPH)
        font1 = head1.style.font
        font1.name = 'Arial Black'
        font1.size = Pt(20)
        font1.bold = True

        head2 = doc.add_paragraph("METROLOGY LABORATORY\n“a SADCAS Accredited Calibration Laboratory, No. CAL-15001”")
        head2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        head2.style = doc.styles.add_style('Style name2', WD_STYLE_TYPE.PARAGRAPH)
        font2 = head2.style.font
        font2.name = 'Arial'
        font2.size = Pt(16)
        font2.bold = True



        para1 = doc.add_paragraph(
            "This certificate of calibration is issued in accordance with section 4 (1) (b) of the "
            "Standards Act No.2 of 2009. The certificate has been issued without any alteration and may "
            "not be published other than in full, except with the prior written approval of the "
            "Director General of Tanzania Bureau of Standards (TBS), P O Box 9524, Dar es Salaam, "
            "Physical Address: Morogoro / Sam Nujoma Road, Ubungo.Tel. +255 22 2450206, Dir. +255 22 "
            "2450298, Fax No. +255 22 2450959")
        para1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        para1.style = doc.styles.add_style('Style name3', WD_STYLE_TYPE.PARAGRAPH)
        font = para1.style.font
        font.name = 'Arial'
        font.size = Pt(10)
        font.bold = True

        para2 = doc.add_paragraph("IDENTIFICATION:")
        para2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        para2.style = doc.styles.add_style('Style name_2', WD_STYLE_TYPE.PARAGRAPH)
        font = para2.style.font
        font.name = 'Arial'
        font.size = Pt(12)
        font.bold = True

        para3 = doc.add_paragraph(f"EQUIPMENT: {self.Equipment} \nMANUFACTURER: {self.Manufacturer} \nIDENTIFICATION "
                                  f"NUMBER: {self.Identification_Number}\n"
                                  f"SERIAL NUMBER:{self.Serial_Number} \nREADABILITY: {self.Readability} "
                                  f"\nLOCATION: {self.Location}\n"
                                  f"TBS JOB NO: {self.Job_num} \nDATE OF CALIBRATION: {self.Calibration_date}\n"
                                  f"DATE OF ISSUE: {self.Issue_date} \nCALIBRATED FOR:\t{self.Customer_Name},"
                                  f"\n\t\t\tP.O.BOX ,{self.Postal_Address},\n\t\t\t{self.Region},{self.Country}.")
        para3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        para3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        para3.paragraph_format.line_spacing = 1.5
        para3.style = doc.styles.add_style('Style name_3', WD_STYLE_TYPE.PARAGRAPH)
        font = para3.style.font
        font.name = 'Arial'
        font.size = Pt(10)
        font.bold = False

        section_f = doc.sections[0]
        footer_1 = section_f.footer
        para_f = footer_1.add_paragraph()
        r = para_f.add_run()
        r.text = f"Calibrated by\t\t\t\t\tChecked by\t\t\t\t\tApproved" \
                 f"by\n\n\n {self.Metrologist}\t\t\t\t\t{self.Technical_Signatory}\t\t" \
                 f"\t\t\t{self.Head_of_department}\nMetrologist\t\t\t\t\tTechnical Signatory" \
                 f"\t\tHead of Metrology Laboratory"
        font = r.font
        font.name = "Arial"
        font.size = Pt(11)
        font.italic = True
        font.bold = False

        para_f_2 = footer_1.add_paragraph()
        para_f_2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        para_f_2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        r_2 = para_f_2.add_run()
        r_2.text = "VALIDITY OF CALIBRATION"
        font = r_2.font
        font.name = "Arial"
        font.size = Pt(8)
        font.italic = False
        font.bold = True
        font.underline = True

        para_f_3 = footer_1.add_paragraph()
        para_f_3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        para_f_3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        r_2 = para_f_3.add_run()
        r_2.text = "The values in this certificate are correct at the time of calibration.  Subsequently the accuracy " \
                   "will depend on such factors as the care exercised in handling and use of the instrument and the " \
                   "frequency of use.  Recalibration should be performed after a period which has been chosen to " \
                   "ensure that the instrument’s accuracy remains within the desired limits. The results applies only " \
                   "to the equipment specified in this document. "
        font = r_2.font
        font.name = "Arial"
        font.size = Pt(8)
        font.italic = False
        font.bold = True
        font.underline = False

        doc.add_section()
        section_2 = doc.sections[1]
        header = section_2.header
        header.is_linked_to_previous = False
        para = header.add_paragraph()
        para.paragraph_format.space_after = 0
        para.paragraph_format.space_before = 0
        r = para.add_run()
        r.add_picture(resource_path("tbs logo.png"))
        r_space_1 = para.add_run("\t\t\t\t\t\t\t\t")
        r_1 = para.add_run()
        r_1.add_picture(resource_path("sadcas.png"))

        para_h = header.add_paragraph(f'\t\t\t\t\tCERTIFICATE NO: {Certificate_No}')
        para_h.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        para_h.style = doc.styles.add_style('Style name_h_2', WD_STYLE_TYPE.PARAGRAPH)
        font = para_h.style.font
        font.name = 'Arial'
        font.size = Pt(12)
        font.bold = True

        r_h = para_h.add_run("\t\t\t\t  Page 2 of 3")
        font = r_h.font
        font.name = 'Arial'
        font.size = Pt(12)
        font.bold = False

        section_f_2 = doc.sections[1]
        footer_2 = section_f_2.footer
        footer_2.is_linked_to_previous = False
        para_f = footer_2.add_paragraph()
        r = para_f.add_run()
        r.text = f"Calibrated by\t\t\t\t\tChecked by\t\t\t\t\t" \
                 f"Approved by\n\n\n {self.Metrologist}\t\t\t\t\t{self.Technical_Signatory}\t\t" \
                 f"\t\t\t{self.Head_of_department}\nMetrologist\t\t\t\t\tTechnical Signatory\t" \
                 f"\tHead of Metrology Laboratory"
        font = r.font
        font.name = "Arial"
        font.size = Pt(11)
        font.italic = True
        font.bold = False

        para4 = doc.add_paragraph("1.\tCALIBRATION CONDITIONS")
        para4.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        para4.paragraph_format.line_spacing = 1.5
        para4.paragraph_format.space_after = 0
        para4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        para4.style = doc.styles.add_style('Style name_4', WD_STYLE_TYPE.PARAGRAPH)
        font = para4.style.font
        font.name = "Arial"
        font.size = Pt(11)
        font.bold = True

        amb_temp = self.Ambient_temperature
        if '.' in amb_temp:
            amb_temp = amb_temp.replace('.', ',')

        rel_humid = self.Relative_humidity
        if '.' in rel_humid:
            rel_humid = rel_humid.replace('.', ',')

        para5 = doc.add_paragraph(f"\tThe calibration was carried out at an ambient "
                                  f"temperature of {amb_temp}°C and "
                                  f"relative humidity of\t{rel_humid}% RH ")
        para5.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        para5.paragraph_format.line_spacing = 1.5
        para5.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        para5.style = doc.styles.add_style('Style name_5', WD_STYLE_TYPE.PARAGRAPH)
        font = para5.style.font
        font.name = "Arial"
        font.size = Pt(11)
        font.bold = False

        para6 = doc.add_paragraph("2.\tEQUIPMENT AND STANDARDS USED")
        para6.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        para6.paragraph_format.line_spacing = 1.5
        para6.paragraph_format.space_after = 0
        para6.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        para6.style = doc.styles.add_style('Style name_6', WD_STYLE_TYPE.PARAGRAPH)
        font = para6.style.font
        font.name = "Arial"
        font.size = Pt(11)
        font.bold = True

        para7 = doc.add_paragraph(f"\tStandard Timer with Identification number {self.Std_timer_id}, with certificate "
                                  f"number {self.Std_timer_cert} Thermo-\thygrometer with serial {self.Thermo_serial},"
                                  f" certificate number {self.Thermo_cert} valid up to {self.Valid}")
        para7.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        para7.paragraph_format.line_spacing = 1.5
        para7.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        para7.style = doc.styles.add_style('Style name_7', WD_STYLE_TYPE.PARAGRAPH)
        font = para7.style.font
        font.name = "Arial"
        font.size = Pt(11)
        font.bold = False

        para8 = doc.add_paragraph("3.\tCALIBRATION PROCEDURE")
        para8.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        para8.paragraph_format.line_spacing = 1.5
        para8.paragraph_format.space_after = 0
        para8.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        para8.style = doc.styles.add_style('Style name_8', WD_STYLE_TYPE.PARAGRAPH)
        font = para8.style.font
        font.name = "Arial"
        font.size = Pt(11)
        font.bold = True

        para9 = doc.add_paragraph(
            f"\tTimer was calibrated by comparison with the standard timer using Method {self.Method_of_calibration} ")
        para9.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        para9.paragraph_format.line_spacing = 1.5
        para9.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        para9.style = doc.styles.add_style('Style name_9', WD_STYLE_TYPE.PARAGRAPH)
        font = para9.style.font
        font.name = "Arial"
        font.size = Pt(11)
        font.bold = False

        para10 = doc.add_paragraph("4\tTRACEABILITY")
        para10.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        para10.paragraph_format.line_spacing = 1.5
        para10.paragraph_format.space_after = 0
        para10.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        para10.style = doc.styles.add_style('Style name_10', WD_STYLE_TYPE.PARAGRAPH)
        font = para10.style.font
        font.name = "Arial"
        font.size = Pt(11)
        font.bold = True

        para11 = doc.add_paragraph(
            f" \tThe results are through regular calibration of the used equipment traceable to the"
            f" Swedish \tNational Laboratory for electrical quantities (RMP 01) at RISE Research Institutes of Sweden.")
        para11.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        para11.paragraph_format.line_spacing = 1.5
        para11.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        para11.style = doc.styles.add_style('Style name_11', WD_STYLE_TYPE.PARAGRAPH)
        font = para11.style.font
        font.name = "Arial"
        font.size = Pt(11)
        font.bold = False
        ron = para11.add_run("    ")

        para12 = doc.add_paragraph("5.\tCALIBRATION RESULTS")
        para12.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        para12.paragraph_format.line_spacing = 1.5
        para12.paragraph_format.space_after = 0
        para12.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        para12.style = doc.styles.add_style('Style name_12', WD_STYLE_TYPE.PARAGRAPH)
        font = para12.style.font
        font.name = "Arial"
        font.size = Pt(11)
        font.bold = True

        if '.' in self.D_1:
            self.D_1 = self.D_1.replace('.', ',')

        dut_1 = self.D_1
        if ',' in self.D_1:
            dut_1 = dut_1.replace(',', '.')

        dut_1 = float(dut_1)

        if '.' in self.S_1:
            self.S_1 = self.S_1.replace('.', ',')

        std_1 = self.S_1
        if ',' in std_1:
            std_1 = std_1.replace(',', '.')

        std_1 = float(std_1)
        c_1 = std_1 - dut_1

        if res == 0.5:
            c_1 = round(c_1)

        elif res == 0.05:
            c_1 = round(c_1, 1)

        elif res == 0.005:
            c_1 = round(c_1, 2)

        if std_1 > dut_1:
            c_1 = f"+{c_1}"
        c_1 = str(c_1)
        if '.' in c_1:
            c_1 = c_1.replace('.', ',')

        if '.' in self.D_2:
            self.D_2 = self.D_2.replace('.', ',')

        dut_2 = self.D_2
        if ',' in self.D_2:
            dut_2 = dut_2.replace(',', '.')

        dut_2 = float(dut_2)

        if '.' in self.S_2:
            self.S_2 = self.S_2.replace('.', ',')

        std_2 = self.S_2
        if ',' in std_2:
            std_2 = std_2.replace(',', '.')

        std_2 = float(std_2)
        c_2 = std_2 - dut_2

        if res == 0.5:
            c_2 = round(c_2)

        elif res == 0.05:
            c_2 = round(c_2, 1)

        elif res == 0.005:
            c_2 = round(c_2, 2)

        if std_2 > dut_2:
            c_2 = f"+{c_2}"
        c_2 = str(c_2)
        if '.' in c_2:
            c_2 = c_2.replace('.', ',')

        if '.' in self.D_3:
            self.D_3 = self.D_3.replace('.', ',')

        dut_3 = self.D_3
        if ',' in self.D_3:
            dut_3 = dut_3.replace(',', '.')

        dut_3 = float(dut_3)

        if '.' in self.S_3:
            self.S_3 = self.S_3.replace('.', ',')

        std_3 = self.S_3
        if ',' in std_3:
            std_3 = std_3.replace(',', '.')

        std_3 = float(std_3)
        c_3 = std_3 - dut_3

        if res == 0.5:
            c_3 = round(c_3)

        elif res == 0.05:
            c_3 = round(c_3, 1)

        elif res == 0.005:
            c_3 = round(c_3, 2)

        if std_3 > dut_3:
            c_3 = f"+{c_3}"
        c_3 = str(c_3)
        if '.' in c_3:
            c_3 = c_3.replace('.', ',')

        if '.' in self.D_4:
            self.D_4 = self.D_4.replace('.', ',')

        dut_4 = self.D_4
        if ',' in self.D_4:
            dut_4 = dut_4.replace(',', '.')

        dut_4 = float(dut_4)

        if '.' in self.S_4:
            self.S_4 = self.S_4.replace('.', ',')

        std_4 = self.S_4
        if ',' in std_4:
            std_4 = std_4.replace(',', '.')

        std_4 = float(std_4)
        c_4 = std_4 - dut_4

        if res == 0.5:
            c_4 = round(c_4)

        elif res == 0.05:
            c_4 = round(c_4, 1)

        elif res == 0.005:
            c_4 = round(c_4, 2)

        if std_4 > dut_4:
            c_4 = f"+{c_4}"
        c_4 = str(c_4)
        if '.' in c_4:
            c_4 = c_4.replace('.', ',')

        if '.' in self.D_5:
            self.D_5 = self.D_5.replace('.', ',')

        dut_5 = self.D_5
        if ',' in self.D_5:
            dut_5 = dut_5.replace(',', '.')

        dut_5 = float(dut_5)

        if '.' in self.S_5:
            self.S_5 = self.S_5.replace('.', ',')

        std_5 = self.S_5
        if ',' in std_5:
            std_5 = std_5.replace(',', '.')

        std_5 = float(std_5)
        c_5 = std_5 - dut_5

        if res == 0.5:
            c_5 = round(c_5)

        elif res == 0.05:
            c_5 = round(c_5, 1)

        elif res == 0.005:
            c_5 = round(c_5, 2)

        if std_5 > dut_5:
            c_5 = f"+{c_5}"
        c_5 = str(c_5)
        if '.' in c_5:
            c_5 = c_5.replace('.', ',')

        uncertainty = str(uncertainty)
        if '.' in uncertainty:
            uncertainty = uncertainty.replace('.', ',')

        data_sheet = [[self.D_1, self.S_1, c_1, uncertainty],
                      [self.D_2, self.S_2, c_2, uncertainty],
                      [self.D_3, self.S_3, c_3, uncertainty],
                      [self.D_4, self.S_4, c_4, uncertainty],
                      [self.D_5, self.S_5, c_5, uncertainty]]

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.allow_autofit = True
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells
        hdr_0 = hdr_cells[0].add_paragraph(f'DUT Time \n interval \n ({self.unit})')
        hdr_0.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        hdr_0.style = doc.styles.add_style('Style name_hdr', WD_STYLE_TYPE.PARAGRAPH)
        font = hdr_0.style.font
        font.name = "Arial"
        font.size = Pt(11)
        font.bold = False
        hdr_1 = hdr_cells[1].add_paragraph(f'Standard Time \n interval \n ({self.unit})')
        font = hdr_1.style.font
        font.name = "Arial"
        font.size = Pt(11)
        font.bold = False
        hdr_1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        hdr_2 = hdr_cells[2].add_paragraph(f'DUT Correction \n ({self.unit})')
        hdr_2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        hdr_3 = hdr_cells[3].add_paragraph(f'Uncertainty of \n Measurement \n ±({self.unit})')
        hdr_3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for a, s, d, f in data_sheet:
            row_cells = table.add_row().cells
            row_0 = row_cells[0].add_paragraph(str(a))
            row_0.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_1 = row_cells[1].add_paragraph(str(s))
            row_1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_2 = row_cells[2].add_paragraph(str(d))
            row_2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_3 = row_cells[3].add_paragraph(str(f))
            row_3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_section()
        section_2 = doc.sections[2]
        header = section_2.header
        header.is_linked_to_previous = False
        para = header.add_paragraph()
        para.paragraph_format.space_after = 0
        para.paragraph_format.space_before = 0
        r = para.add_run()
        r.add_picture(resource_path("tbs logo.png"))
        r_space_1 = para.add_run("\t\t\t\t\t\t\t\t")
        r_1 = para.add_run()
        r_1.add_picture(resource_path("sadcas.png"))

        para_h = header.add_paragraph(f'\t\t\t\t\tCERTIFICATE NO: {Certificate_No}')
        para_h.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        para_h.style = doc.styles.add_style('Style name_h_3', WD_STYLE_TYPE.PARAGRAPH)
        font = para_h.style.font
        font.name = 'Arial'
        font.size = Pt(12)
        font.bold = True

        r_h = para_h.add_run("\t\t\t\t  Page 3 of 3")
        font = r_h.font
        font.name = 'Arial'
        font.size = Pt(12)
        font.bold = False

        para18 = doc.add_paragraph("6.0 UNCERTAINTY")
        para18.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        para18.paragraph_format.line_spacing = 1.5
        para18.paragraph_format.space_after = 0
        para18.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        para18.style = doc.styles.add_style('Style name_18', WD_STYLE_TYPE.PARAGRAPH)
        font = para18.style.font
        font.name = "Arial"
        font.size = Pt(11)
        font.bold = True

        para19 = doc.add_paragraph(
            f'    6.1\tThe reported uncertainties of measurement were calculated in accordance with the BIPM,'
            f'IEC, ISO, \tIUPAP, OIML document entitled “Guide to the Expression of Uncertainty in'
            f' \tMeasurement“(International Organisation for Standardisation, Geneva, Switzerland, 2008).')

        para19.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        para19.paragraph_format.line_spacing = 1.5
        para19.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        para19.style = doc.styles.add_style('Style name_19', WD_STYLE_TYPE.PARAGRAPH)
        font = para19.style.font
        font.name = "Arial"
        font.size = Pt(11)
        font.bold = False

        para20 = doc.add_paragraph(
            f'    6.2 \tThe reported expanded uncertainty of measurement is stated as the standard uncertainty of'
            f' \tmeasurement multiplied by a coverage factor of k=2, which for a normal distribution'
            f' approximates a \tlevel of confidence of 95%.')

        para20.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        para20.paragraph_format.line_spacing = 1.5
        para20.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        para20.style = doc.styles.add_style('Style name_20', WD_STYLE_TYPE.PARAGRAPH)
        font = para20.style.font
        font.name = "Arial"
        font.size = Pt(11)
        font.bold = False

        para21 = doc.add_paragraph("7.0 REMARKS")
        para21.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        para21.paragraph_format.line_spacing = 1.5
        para21.paragraph_format.space_after = 0
        para21.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        para21.style = doc.styles.add_style('Style name_21', WD_STYLE_TYPE.PARAGRAPH)
        font = para21.style.font
        font.name = "Arial"
        font.size = Pt(11)
        font.bold = True

        para22 = doc.add_paragraph(f'    7.1\t{self.Remarks}')

        para22.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        para22.paragraph_format.line_spacing = 1.5
        para22.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        para22.style = doc.styles.add_style('Style name_22', WD_STYLE_TYPE.PARAGRAPH)
        font = para22.style.font
        font.name = "Arial"
        font.size = Pt(11)
        font.bold = False

        para25 = doc.add_paragraph("\n\n\n“END OF CERTIFICATE”")
        para25.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        para25.paragraph_format.line_spacing = 1.5
        para25.paragraph_format.space_after = 0
        para25.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        para25.style = doc.styles.add_style('Style name_25', WD_STYLE_TYPE.PARAGRAPH)
        font = para25.style.font
        font.name = "Arial"
        font.size = Pt(11)
        font.bold = True

        doc.save(resource_path(f'Certificates\{Certificate_No}.docx'))


kvfile = Builder.load_string("""#: import Factory kivy.factory.Factory
MainLayout:

<MyPopup@Popup>:
    auto_dismiss: False
    title: "Generator Info - Lordrick Apps"
    size_hint: 0.5,0.3
    pos_hint: {"center_x": 0.5}
    BoxLayout:
        orientation: "vertical"
        Label:
            text: "Congratulations!!!Your Certificate and Uncertainty Budget have been generated"
            color: '#00FFCE'
        Label:
            text: "Check Certificates Folder"
            color: '#00FFCE'
        Button:
            text: "Ok, Thanks!!"
            color: '#00FFCE'
            size_hint: 0.5,0.8
            pos_hint: {"center_x": 0.5}
            on_release: root.dismiss()
<MainLayout>:
    do_default_tab: False
    tab_width: "200dp"

    TabbedPanelItem:
        text: "Customer Details"

        BoxLayout:
            orientation: "vertical"
            Label:
                text: "Enter Customer Details here"
                size_hint: 0.5,0.2
                pos_hint: {"center_x": 0.5}

            GridLayout:
                cols: 4
                Label:
                    text: "Certificate Number:"

                TextInput:
                    text: "9999"
                    on_focus: "9999"
                    multiline: False
                    on_text_validate: root.cert_numb(self)

                Label:
                    text: "Equipment:"

                TextInput:
                    text: "TIMER"
                    on_focus: "TIMER"
                    multiline: False
                    on_text_validate: root.equipment(self)

                Label:
                    text: "Identification Number:"

                TextInput:
                    text: "TBS/LAB/01"
                    on_focus: "TBS/LAB/01"
                    multiline: False
                    on_text_validate: root.identification_number(self)

                Label:
                    text: "Manufacturer:"

                TextInput:
                    text: " - "
                    on_focus: " - "
                    multiline: False
                    on_text_validate: root.manufacturer(self)

                Label:
                    text: "Serial Number:"

                TextInput:
                    text: " 9999 "
                    on_focus: " 9999 "
                    multiline: False
                    on_text_validate: root.serial_number(self)

                Label:
                    text: "Readability"

                TextInput:
                    text: "0,01s"
                    on_focus: "0,01s"
                    multiline: False
                    on_text_validate: root.readability(self)

                Label:
                    text: "Location:"

                TextInput:
                    text: " TBS TIME AND FREQUENCY LAB UBUNGO "
                    on_focus: " TBS TIME AND FREQUENCY LAB UBUNGO "
                    multiline: False
                    on_text_validate: root.location(self)

                Label:
                    text: "TBS Job No.:"

                TextInput:
                    text: " 2021 - 999"
                    on_focus: " 2021 - 999 "
                    multiline: False
                    on_text_validate: root.job_num(self)

                Label:
                    text: "Date of Calibration:"

                TextInput:
                    text: root.d1
                    on_focus: root.d1
                    multiline: False
                    on_text_validate: root.calibration_date(self)

                Label:
                    text: "Date of Issue:"

                TextInput:
                    text: root.d1
                    on_focus: root.d1
                    multiline: False
                    on_text_validate: root.issue_date(self)

                Label:
                    text: "Customer Name:"

                TextInput:
                    text: "Lordrick Apps"
                    on_focus: "Lordrick Apps"
                    multiline: False
                    on_text_validate: root.customer_name(self)

                Label:
                    text: "P. O. Box:"

                TextInput:
                    text: "7713"
                    on_focus: "7713"
                    multiline: False
                    on_text_validate: root.postal_address(self)

                Label:
                    text: "Region"

                TextInput:
                    text: "Dodoma"
                    on_focus: "Dodoma"
                    multiline: False
                    on_text_validate: root.region(self)

                Label:
                    text: "Country"

                TextInput:
                    text: "Tanzania"
                    on_focus: "Tanzania"
                    multiline: False
                    on_text_validate: root.country(self)

            Label:
                text: "Your Current Data"
                size_hint: 0.5,0.1
                pos_hint: {"center_x": 0.5}

            GridLayout:
                rows: 7

                Label:
                    text: "The Certificate Number Entered was " + root.certificate_number

                Label:
                    text: "The Equipment name Entered was " + root.Equipment

                Label:
                    text: "The Identification Number Entered was " + root.Identification_Number

                Label:
                    text: "The Manufacturer name Entered was " + root.Manufacturer

                Label:
                    text: "The Serial Number Entered was " + root.Serial_Number

                Label:
                    text: "The readability Entered was " + root.Readability

                Label:
                    text: "The Location Entered was " + root.Location

                Label:
                    text: "The TBS job no. Entered was " + root.Job_num

                Label:
                    text: "The Calibration date Entered was " + root.Calibration_date

                Label:
                    text: "The Issue date Entered was " + root.Issue_date

                Label:
                    text: "The Customer name Entered was " + root.Customer_Name

                Label:
                    text: "The Postal Address Entered was P. O. Box " + root.Postal_Address

                Label:
                    text: "The region Entered was " + root.Region

                Label:
                    text: "The country Entered was " + root.Country

    TabbedPanelItem:
        text: "Metrological Details"
        BoxLayout:

            BoxLayout:
                orientation: "vertical"
                Label:
                    text: "Enter the Metrological details Below"
                    size_hint: 0.5,0.2
                    pos_hint: {"center_x": 0.5}

                GridLayout:
                    cols: 4

                    Label:
                        text: "Metrologist"

                    TextInput:
                        text: "J. Yarrot"
                        on_focus: "J. Yarrot"
                        multiline: False
                        on_text_validate: root.metrologist(self)

                    Label:
                        text: "Technical Signatory"

                    TextInput:
                        text: "K. Shemhillu"
                        on_focus: "K. Shemhillu"
                        multiline: False
                        on_text_validate: root.technical_signatory(self)

                    Label:
                        text: "Head of Metrology"

                    TextInput:
                        text: "J. Mahilla"
                        on_focus: "J. Mahilla"
                        multiline: False
                        on_text_validate: root.head_of_department(self)

                    Label:
                        text: "Standard Timer Identification Number"

                    TextInput:
                        text: "TBS 8-381 "
                        on_focus: "TBS 8-381"
                        multiline: False
                        on_text_validate: root.std_timer_id(self)

                    Label:
                        text: "Standard Timer Certificate Number:"

                    TextInput:
                        text: "58469"
                        on_focus: "58469"
                        multiline: False
                        on_text_validate: root.std_timer_cert(self)

                    Label:
                        text: "Thermo-hygrometer serial number:"

                    TextInput:
                        text: "MAC:98:8B:AD:20:C2:9F"
                        on_focus: "MAC:98:8B:AD:20:C2:9F"
                        multiline: False
                        on_text_validate: root.thermo_serial(self)

                    Label:
                        text: "Thermo-hygrometer cert number"

                    TextInput:
                        text: "45326"
                        on_focus: "45326"
                        multiline: False
                        on_text_validate: root.thermo_cert(self)

                    Label:
                        text: "Valid up to:"

                    TextInput:
                        text: "Dec 2021"
                        on_focus: "Dec 2021"
                        multiline: False
                        on_text_validate: root.valid(self)

                    Label:
                        text: "Method of Calibration"

                    TextInput:
                        text: "MET - TF - 03"
                        on_focus: "MET - TF - 03"
                        multiline: False
                        on_text_validate: root.method_of_calibration(self)

                    Label:
                        text: "Resolution of UUT"

                    TextInput:
                        text: "0.05"
                        on_focus: "0.05"
                        multiline: False
                        on_text_validate: root.resolution(self)

                Label:
                    text: "Your Current Data"
                    size_hint: 0.5,0.1
                    pos_hint: {"center_x": 0.5}

                GridLayout:
                    rows: 5

                    Label:
                        text: "The Metrologist name Entered was " + root.Metrologist

                    Label:
                        text: "The Technical Signatory Entered was " + root.Technical_Signatory

                    Label:
                        text: "The Head of department name Entered was " + root.Head_of_department

                    Label:
                        text: "The Standard timer identification number Entered was " + root.Std_timer_id

                    Label:
                        text: "The Standard Timer certificate number Entered was " + root.Std_timer_cert

                    Label:
                        text: "The Thermo-hygrometer serial number  Entered was " + root.Thermo_serial

                    Label:
                        text: "The Thermo-hygrometer certificate number Entered was " + root.Thermo_cert

                    Label:
                        text: "The Valid up to  Entered was " + root.Valid

                    Label:
                        text: "The method of calibration Entered was " + root.Method_of_calibration

                    Label:
                        text: "The resolution Entered was " + root.Resolution

    TabbedPanelItem:
        text: "Data Collected"
        BoxLayout:
            orientation: "vertical"
            GridLayout:
                size_hint: 1,0.1
                cols: 4

                Label:
                    text: "Ambient Temperature"

                TextInput:
                    text: "22,00"
                    on_focus: "22,00"
                    multiline: False
                    on_text_validate: root.ambient_temperature(self)

                Label:
                    text: "Relative Humidity"

                TextInput:
                    text: "50,0"
                    on_focus: "50,0"
                    multiline: False
                    on_text_validate: root.relative_humidity(self)

            GridLayout:
                cols: 2
                size_hint: 1,0.2
                Label:
                    text: "The Ambient Temperature entered was "+ root.Ambient_temperature

                Label:
                    text: "The Relative Humidity entered was "+ root.Relative_humidity


            BoxLayout:
                BoxLayout:
                    orientation: "vertical"
                    GridLayout:
                        size_hint: 1,0.1
                        cols: 4

                        Label:
                            text: "seconds"

                        CheckBox:
                            group: "units"
                            on_active: root.on_click_seconds(self,self.active,"seconds")

                        Label:
                            text: "minutes"

                        CheckBox:
                            group: "units"
                            on_active: root.on_click_minutes(self,self.active,"minutes")
                    Label:
                        text: "Enter the data collected"
                        size_hint: 0.5,0.1
                        pos_hint: {"center_x": 0.5}

                    GridLayout:
                        cols: 2
                        Label:
                            text: "UUT readings"

                        Label:
                            text: "Standard readings"

                        TextInput:
                            text: "60"
                            on_focus: "60"
                            multiline: False
                            on_text_validate: root.d_1(self)

                        TextInput:
                            text: "60"
                            on_focus: "60"
                            multiline: False
                            on_text_validate: root.s_1(self)

                        TextInput:
                            text: "300"
                            on_focus: "300"
                            multiline: False
                            on_text_validate: root.d_2(self)

                        TextInput:
                            text: "300"
                            on_focus: "300"
                            multiline: False
                            on_text_validate: root.s_2(self)

                        TextInput:
                            text: "600"
                            on_focus: "600"
                            multiline: False
                            on_text_validate: root.d_3(self)

                        TextInput:
                            text: "600"
                            on_focus: "600"
                            multiline: False
                            on_text_validate: root.s_3(self)

                        TextInput:
                            text: "900"
                            on_focus: "900"
                            multiline: False
                            on_text_validate: root.d_4(self)

                        TextInput:
                            text: "900"
                            on_focus: "900"
                            multiline: False
                            on_text_validate: root.s_4(self)

                        TextInput:
                            text: "1200"
                            on_focus: "1200"
                            multiline: False
                            on_text_validate: root.d_5(self)

                        TextInput:
                            text: "1200"
                            on_focus: "1200"
                            multiline: False
                            on_text_validate: root.s_5(self)

                BoxLayout:
                    orientation: "vertical"
                    Label:
                        text: "Your Current Data"
                        size_hint: 0.5,0.1
                        pos_hint: {"center_x": 0.5}

                    GridLayout:
                        cols: 2
                        Label:
                            text: "UUT readings (" + root.unit + ")"

                        Label:
                            text: "Standard readings (" + root.unit + ")"

                        Label:
                            text: root.D_1

                        Label:
                            text: root.S_1

                        Label:
                            text: root.D_2

                        Label:
                            text: root.S_2

                        Label:
                            text: root.D_3

                        Label:
                            text: root.S_3

                        Label:
                            text: root.D_4

                        Label:
                            text: root.S_4

                        Label:
                            text: root.D_5

                        Label:
                            text: root.S_5
            GridLayout:
                size_hint: 1,0.4
                cols:2 
                Label:
                    text: "Enter Your remarks Below"
                    
                Label:
                    text: "Your remarks entered are"
                    
                TextInput:
                    text: "DUT Means device under test"
                    on_focus: "DUT Means device under test"
                    multiline: False
                    on_text_validate: root.remarks(self)
                    
                Label:
                    text: root.Remarks

            Button:
                text: "Generate Certificate"
                color: '#00FFCE'
                font_name: 'Lemonada'
                font_size: '22dp'
                size_hint: 0.25,0.25
                pos_hint: {"center_x": 0.5}
                on_press: root.certificate_generator()
                on_release: Factory.MyPopup().open()
                """)


class TIMER_CERTIFICATE_GENERATORApp(App):
    def build(self):
        return MainLayout()


TIMER_CERTIFICATE_GENERATORApp().run()
